import re
import pdfplumber
import pytesseract
from pdf2image import convert_from_path
from PIL import Image
import asyncio
from pathlib import Path
from typing import Dict, List, Optional, Tuple
import io
import math
from backend.intelligence.pdf_extractor import TenderPDFExtractor

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

class DocumentExtractor:
    """
    HONEST EXTRACTION - Multi-Sector Tender Intelligence
    Zero placeholders. Real estimates only when explicitly marked.
    """
    
    def __init__(self):
        self.extraction_confidence = {}
        self.estimated_values = {}
        self.pdf_intel = TenderPDFExtractor()
    
    async def extract(self, file_path: str) -> Dict:
        """Extract with sector-specific intelligence"""
        path = Path(file_path)
        ext = path.suffix.lower()
        
        # Extract raw text
        if ext == '.pdf':
            text = await self._extract_pdf(path)
        elif ext in ['.docx', '.doc']:
            text = await self._extract_docx(path)
        else:
            text = await self._extract_text(path)
        
        # KEEP FULL TEXT for comprehensive search
        full_text = text
        text_upper = text.upper()
        
        # DETECT SECTOR
        sector_data = self._detect_sector_and_type(text_upper)
        sector = sector_data['sector']
        sub_type = sector_data['sub_type']
        
        # Extract duration with context
        duration_data = self._extract_duration_intelligent(text_upper, sector)
        
        # Extract or estimate workers using FULL TEXT
        worker_data = self._estimate_workers_by_sector(text_upper, full_text, sector, duration_data['months'])
        
        # Extract or estimate area/scope
        scope_data = self._estimate_scope_by_sector(text_upper, full_text, sector)
        
        # Extract location
        location = self._extract_location(text_upper)
        
        # USE PRODUCTION PDF EXTRACTOR
        intel_data = self.pdf_intel.process_text(full_text)
        estimated_value = intel_data['estimated_value']
        
        # Build result
        total_workers = worker_data['total']
        supervisors = worker_data.get('supervisors', 0)
        duration_months = duration_data.get('months', 12)
        area_sqm = scope_data.get('area', 1000)
        shifts = self._extract_shifts(text_upper)
        hours = self._extract_hours(text_upper, sector)
        working_days = self._extract_working_days(text_upper)
        
        result = {
            'reference': self._extract_reference(text_upper) or f"REF-{path.stem[:15]}",
            'closing_date': self._extract_closing_date(full_text),
            'sector': sector,
            'sub_type': sub_type,
            'location': location,
            'estimated_value': estimated_value,
            
            'duration': duration_data,
            
            'scope': {
                'area_sqm': area_sqm,
                'area_note': scope_data.get('note', 'Extracted from document'),
                'description': self._extract_scope_description(full_text, sector),
                'is_emergency': 'EMERGENCY' in text_upper or 'URGENT' in text_upper,
                'is_breakdown': 'BREAKDOWN' in text_upper or 'REPAIR' in text_upper
            },
            
            'workforce': {
                'total_workers': total_workers,
                'skilled_workers': worker_data.get('skilled', 0),
                'unskilled_workers': worker_data.get('unskilled', 0),
                'supervisors': supervisors,
                'estimation_method': worker_data.get('method', 'Extracted'),
                'confidence': worker_data.get('confidence', 'High')
            },
            
            'requirements': {
                'shifts_per_day': shifts,
                'hours_per_day': hours,
                'working_days_week': working_days,
                'equipment_required': self._extract_equipment_by_sector(text_upper, sector),
                'materials_required': self._extract_materials_by_sector(text_upper, sector),
                'certifications_required': self._extract_certifications(text_upper, sector),
                # Backward-compat fields for document_generator, debate_council, index.html
                'workers': total_workers,
                'supervisors': supervisors,
                'area_sqm': area_sqm,
                'shifts': shifts,
                'total_staff': total_workers + supervisors
            },
            
            'client': {
                'type': self._extract_client_type(text_upper),
                'name': self._extract_client_name(full_text)
            },
            
            '_extraction_notes': {
                'sector_detected': sector,
                'worker_estimated': worker_data.get('estimated', False),
                'duration_estimated': duration_data.get('estimated', False),
                'area_estimated': scope_data.get('estimated', False),
                'raw_confidence': self._calculate_overall_confidence(worker_data, duration_data, scope_data)
            },
            
            'raw_text': full_text[:5000],
            
            # === BACKWARD-COMPAT TOP-LEVEL FIELDS ===
            # These duplicate nested data so downstream code using the old flat format still works
            'tender_type': sector,
            'duration_months': duration_months
        }
        
        return result
    
    def _detect_sector_and_type(self, text: str) -> Dict:
        """Detect sector with weighted scoring - FIXED for cleaning detection"""
        
        # CRITICAL FIX #1: Check for exact phrase "CLEANING SERVICES" first (high priority)
        if 'CLEANING SERVICES' in text:
            return {'sector': 'cleaning', 'sub_type': 'commercial', 'confidence': 15}
        
        # Also check common cleaning patterns
        cleaning_phrases = ['OFFICE CLEANING', 'CLEANING CONTRACT', 'JANITORIAL SERVICES', 
                           'HYGIENE SERVICES', 'CLEANING AND', 'CLEANING OF']
        for phrase in cleaning_phrases:
            if phrase in text:
                return {'sector': 'cleaning', 'sub_type': 'standard', 'confidence': 12}
        
        # CONSTRUCTION
        construction_terms = {
            'construction': 5, 'building': 5, 'civil': 5, 'structural': 4,
            'concrete': 4, 'brickwork': 4, 'plastering': 4, 'roofing': 4,
            'excavation': 4, 'foundation': 4, 'renovation': 3, 'refurbishment': 3,
            'plumbing': 3, 'carpentry': 3, 'painting': 3, 'paving': 3,
            'asphalt': 4, 'tar': 3, 'road': 3, 'bridge': 4,
            'demolition': 3, 'earthworks': 4, 'steel': 3, 'formwork': 4,
            'scaffolding': 3, 'drainage': 3, 'sewer': 3, 'waterproofing': 3
        }
        scores = {'construction': sum(weight for term, weight in construction_terms.items() if term in text)}
        
        # ELECTRICAL
        electrical_terms = {
            'electrical': 5, 'electrician': 5, 'wiring': 4, 'cable': 3,
            'distribution board': 5, 'substation': 4, 'high voltage': 5,
            'solar': 4, 'generator': 3, 'reticulation': 3, 'transformer': 4
        }
        scores['electrical'] = sum(weight for term, weight in electrical_terms.items() if term in text)
        
        # CLEANING (individual terms - lower weight than phrases above)
        cleaning_terms = {
            'cleaning': 3, 'cleaner': 3, 'hygiene': 4, 'sanitiz': 4,
            'janitorial': 4, 'housekeeping': 3, 'deep clean': 4
        }
        scores['cleaning'] = sum(weight for term, weight in cleaning_terms.items() if term in text)
        
        # SECURITY
        security_terms = {
            'security': 5, 'guarding': 5, 'cctv': 4, 'surveillance': 4,
            'access control': 5, 'alarm': 3, 'patrol': 4, 'psira': 5
        }
        scores['security'] = sum(weight for term, weight in security_terms.items() if term in text)
        
        # GARDENING
        garden_terms = {
            'garden': 5, 'landscaping': 5, 'grass': 3, 'lawn': 3,
            'tree': 3, 'felling': 4, 'pruning': 4, 'irrigation': 4,
            'grounds': 4, 'horticulture': 5
        }
        scores['gardening'] = sum(weight for term, weight in garden_terms.items() if term in text)
        
        # IT
        it_terms = {
            'software': 5, 'programming': 5, 'development': 4, 'it services': 5,
            'network': 4, 'server': 4, 'cloud': 4, 'cyber': 4,
            'database': 4, 'application': 3, 'helpdesk': 4,
            'coding': 5, 'java': 4, 'python': 4, 'sql': 4
        }
        scores['it_services'] = sum(weight for term, weight in it_terms.items() if term in text)
        
        # MAINTENANCE
        maintenance_terms = {
            'maintenance': 4, 'repair': 3, 'service': 2, 'breakdown': 4
        }
        scores['maintenance'] = sum(weight for term, weight in maintenance_terms.items() if term in text)
        
        # SUPPLY
        supply_terms = {
            'supply': 5, 'deliver': 4, 'procure': 4, 'materials': 3,
            'equipment': 3, 'goods': 3, 'tools': 3
        }
        scores['supply'] = sum(weight for term, weight in supply_terms.items() if term in text)
        
        # Determine winner
        best_sector = max(scores, key=scores.get)
        
        # Handle low confidence
        if scores[best_sector] < 5:
            best_sector = 'general'
            confidence = 0
        else:
            confidence = scores[best_sector]
        
        # Detect sub-type
        sub_type = 'standard'
        if 'EMERGENCY' in text or 'URGENT' in text:
            sub_type = 'emergency'
        elif 'RESIDENTIAL' in text or 'HOUSING' in text:
            sub_type = 'residential'
        elif 'INDUSTRIAL' in text or 'FACTORY' in text:
            sub_type = 'industrial'
        elif 'COMMERCIAL' in text or 'OFFICE' in text:
            sub_type = 'commercial'
        elif 'REPAIR' in text or 'BREAKDOWN' in text:
            sub_type = 'repair'
        
        return {'sector': best_sector, 'sub_type': sub_type, 'confidence': confidence}
    
    def _estimate_workers_by_sector(self, text: str, raw_text: str, sector: str, duration_months: int) -> Dict:
        """HONEST ESTIMATION: Search full document for workers"""
        
        # Search PATTERNS in priority order
        patterns = [
            r'(\d{1,3})\s*CLEANERS?',
            r'(\d{1,3})\s*WORKERS?',
            r'(\d{1,3})\s*STAFF',
            r'(\d{1,3})\s*PERSONNEL',
            r'(\d{1,3})\s*EMPLOYEES?',
            r'(\d{1,3})\s*OPERATIVES?',
            r'(\d{1,3})\s*LABOURERS?',
            r'(\d{1,3})\s*PEOPLE',
            r'TEAM\s*OF\s*(\d{1,3})',
            r'NO[.:]?\s*OF\s*CLEANERS?\s*:?\s*(\d{1,3})',
            r'NO[.:]?\s*OF\s*WORKERS?\s*:?\s*(\d{1,3})',
            r'QUANTITY[:\s]+(\d{1,3})\s*(?:CLEANERS?|WORKERS?)',
            r'(\d{1,3})\s*DAY\s*SHIFT',
            r'(\d{1,3})\s*NIGHT\s*SHIFT'
        ]
        
        # Search in FULL raw_text (not just first 1000 chars)
        for pattern in patterns:
            matches = re.findall(pattern, raw_text.upper())
            if matches:
                try:
                    # Handle both direct match and group match
                    match = matches[0]
                    workers = int(match) if isinstance(match, str) else int(match[0])
                    if 1 <= workers <= 500:
                        return {
                            'total': workers,
                            'method': 'Extracted from document',
                            'estimated': False,
                            'confidence': 'High',
                            'note': f'Found: {workers} workers'
                        }
                except:
                    continue
        
        # ESTIMATION LOGIC (Clearly marked as estimated)
        estimation = {'estimated': True, 'method': 'Calculated based on sector norms'}
        
        if sector == 'cleaning':
            area_match = re.search(r'(\d{1,6})\s*(?:M2|M²|SQM|SQUARE)', raw_text.upper())
            if area_match:
                area = int(area_match.group(1))
                workers = max(2, area // 450)
                estimation.update({
                    'total': workers,
                    'unskilled': workers,
                    'supervisors': max(1, workers // 10),
                    'note': f'Estimated: {area}m2 / 450m2 per worker = {workers} workers',
                    'confidence': 'Medium'
                })
            else:
                estimation.update({
                    'total': 5,
                    'unskilled': 5,
                    'supervisors': 1,
                    'note': 'Estimated: Standard office cleaning team (5 workers + 1 supervisor)',
                    'confidence': 'Low'
                })
        
        elif sector == 'construction':
            if 'ROAD' in raw_text.upper() or 'ASPHALT' in raw_text.upper():
                estimation.update({
                    'total': 15,
                    'skilled': 8,
                    'unskilled': 6,
                    'supervisors': 1,
                    'note': 'Estimated: Road construction crew (15 workers)',
                    'confidence': 'Medium'
                })
            else:
                estimation.update({
                    'total': 10,
                    'skilled': 6,
                    'unskilled': 3,
                    'supervisors': 1,
                    'note': 'Estimated: General construction crew',
                    'confidence': 'Low'
                })
        
        elif sector == 'electrical':
            estimation.update({
                'total': 4,
                'skilled': 3,
                'unskilled': 0,
                'supervisors': 1,
                'note': 'Estimated: Electrical team (4 workers)',
                'confidence': 'Medium'
            })
        
        elif sector == 'security':
            shifts = 3 if '24 HOUR' in raw_text.upper() else 2
            posts = 2
            total = shifts * posts
            estimation.update({
                'total': total,
                'skilled': total,
                'supervisors': max(1, total // 8),
                'note': f'Estimated: {posts} posts x {shifts} shifts = {total} guards',
                'confidence': 'Medium'
            })
        
        elif sector == 'gardening':
            area_match = re.search(r'(\d{1,6})\s*(?:M2|M²|SQM)', raw_text.upper())
            if area_match:
                area = int(area_match.group(1))
                workers = max(2, area // 2500)
                estimation.update({
                    'total': workers,
                    'unskilled': workers,
                    'note': f'Estimated: {area}m2 / 2500m2 = {workers} workers',
                    'confidence': 'Medium'
                })
            else:
                estimation.update({
                    'total': 3,
                    'unskilled': 3,
                    'note': 'Estimated: Standard gardening team',
                    'confidence': 'Low'
                })
        
        else:
            estimation.update({
                'total': 5,
                'note': 'Estimated: General workforce (5 workers) - VERIFY MANUALLY',
                'confidence': 'Low'
            })
        
        return estimation
    
    def _extract_duration_intelligent(self, text: str, sector: str) -> Dict:
        """FIXED: Look for contract period, not RFQ validity"""
        
        # PRIORITY 1: Look for "period of twelve (12) months" or similar
        contract_patterns = [
            r'period\s+of\s+(\w+)\s+\((\d{1,2})\)\s*months?',  # "twelve (12) months"
            r'period\s+of\s+(\d{1,2})\s*months?',  # "12 months"
            r'agreement.*?for\s+(\d{1,2})\s*months',  # "agreement for 12 months"
            r'duration.*?(\d{1,2})\s*months',  # "duration: 12 months"
            r'contract.*?(\d{1,2})\s*months',  # "contract period 12 months"
            r'for\s+a\s+period\s+of\s+(\d{1,2})\s*months',  # "for a period of 12 months"
            r'twelve\s*\(12\)\s*months',  # "twelve (12) months"
            r'\(12\)\s*months',  # "(12) months"
        ]
        
        for pattern in contract_patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                try:
                    months = int(match.group(2)) if len(match.groups()) > 1 and match.group(2) else int(match.group(1))
                except:
                    continue
                
                return {
                    'value': months,
                    'unit': 'months',
                    'months': months,
                    'display': f'{months} months',
                    'estimated': False,
                    'note': 'Extracted from contract period'
                }
        
        # Check for written "TWELVE"
        if 'TWELVE' in text and 'MONTH' in text:
            return {
                'value': 12,
                'unit': 'months',
                'months': 12,
                'display': '12 months',
                'estimated': False,
                'note': 'Extracted: "twelve months"'
            }
        
        # Years
        year_match = re.search(r'(\d)\s*YEAR', text)
        if year_match:
            years = int(year_match.group(1))
            return {
                'value': years,
                'unit': 'years',
                'months': years * 12,
                'display': f'{years} years',
                'estimated': False
            }
        
        # Emergency days
        if 'EMERGENCY' in text:
            day_match = re.search(r'(\d{1,3})\s*DAYS?', text)
            if day_match:
                days = int(day_match.group(1))
                return {
                    'value': days,
                    'unit': 'days',
                    'months': days / 30,
                    'display': f'{days} days',
                    'estimated': False,
                    'is_emergency': True
                }
        
        # DEFAULT ESTIMATES
        result = {'estimated': True}
        if sector == 'cleaning':
            result.update({
                'value': 12,
                'unit': 'months',
                'months': 12,
                'display': '12 months (estimated)',
                'note': 'Standard cleaning contract duration'
            })
        elif sector == 'construction':
            result.update({
                'value': 6,
                'unit': 'months',
                'months': 6,
                'display': '6 months (estimated)',
                'note': 'Construction project duration'
            })
        else:
            result.update({
                'value': 12,
                'unit': 'months',
                'months': 12,
                'display': '12 months (default)',
                'note': 'Standard contract duration'
            })
        
        return result
    
    def _estimate_scope_by_sector(self, text: str, raw_text: str, sector: str) -> Dict:
        """Estimate scope"""
        area_match = re.search(r'(\d{1,6}(?:,\d{3})?)\s*(?:M2|M²|SQM|SQUARE)', text)
        if area_match:
            return {
                'area': int(area_match.group(1).replace(',', '')),
                'estimated': False,
                'note': 'Extracted from document'
            }
        
        estimation = {'estimated': True}
        
        if sector == 'cleaning':
            offices = re.search(r'(\d{1,3})\s*(?:OFFICES|ROOMS|FLOORS)', text)
            if offices:
                num = int(offices.group(1))
                est_area = num * 50
                estimation.update({
                    'area': est_area,
                    'note': f'Estimated: {num} offices x 50m2 = {est_area}m2'
                })
            else:
                estimation.update({
                    'area': 1000,
                    'note': 'Estimated: Standard office (1000m2)'
                })
        
        elif sector == 'construction':
            houses = re.search(r'(\d{1,3})\s*(?:HOUSE|HOUSES|UNITS)', text)
            if houses:
                num = int(houses.group(1))
                est_area = num * 150
                estimation.update({
                    'area': est_area,
                    'note': f'Estimated: {num} units x 150m2 = {est_area}m2'
                })
            else:
                estimation.update({
                    'area': 5000,
                    'note': 'Estimated: Medium construction project'
                })
        
        else:
            estimation.update({
                'area': 1000,
                'note': 'Estimated: Default scope (1000m2)'
            })
        
        return estimation
    
    def _extract_equipment_by_sector(self, text: str, sector: str) -> List[str]:
        text_lower = text.lower()
        equipment = []
        
        if sector == 'cleaning':
            items = ['scrubber', 'vacuum', 'polisher', 'pressure washer', 'carpet cleaner']
            for item in items:
                if item in text_lower:
                    equipment.append(item)
        
        elif sector == 'construction':
            items = ['excavator', 'tipper', 'mixer', 'scaffolding', 'compactor', 'crane']
            for item in items:
                if item in text_lower:
                    equipment.append(item)
        
        return equipment if equipment else ['Standard equipment per sector']
    
    def _extract_materials_by_sector(self, text: str, sector: str) -> List[str]:
        text_lower = text.lower()
        materials = []
        
        if sector == 'cleaning':
            items = ['detergent', 'disinfectant', 'cloths', 'paper towels']
        elif sector == 'construction':
            items = ['cement', 'bricks', 'steel', 'sand']
        elif sector == 'electrical':
            items = ['cable', 'distribution board', 'breakers']
        else:
            return ['Materials per scope']
        
        for item in items:
            if item in text_lower:
                materials.append(item)
        
        return materials if materials else ['To be specified']
    
    def _extract_certifications(self, text: str, sector: str) -> List[str]:
        certs = []
        if sector == 'electrical':
            certs.append('Master Electrician Certificate')
        elif sector == 'security':
            certs.append('PSIRA Registration')
        elif sector == 'cleaning':
            certs.append('Cleaning Industry Registration')
        elif sector == 'construction':
            certs.append('CIDB Registration')
        return certs if certs else ['Standard registrations']
    
    def _extract_shifts(self, text: str) -> int:
        if '3 SHIFT' in text or 'THREE SHIFT' in text or '24 HOUR' in text:
            return 3
        if '2 SHIFT' in text or 'TWO SHIFT' in text:
            return 2
        return 1
    
    def _extract_hours(self, text: str, sector: str) -> int:
        match = re.search(r'(\d{1,2})\s*HOURS?', text)
        if match:
            hours = int(match.group(1))
            if 4 <= hours <= 12:
                return hours
        if sector == 'security':
            return 12
        return 8
    
    def _extract_working_days(self, text: str) -> int:
        if '7 DAYS' in text:
            return 7
        if '6 DAYS' in text:
            return 6
        if '5 DAYS' in text:
            return 5
        return 5
    
    def _extract_location(self, text: str) -> str:
        cities = ['johannesburg', 'pretoria', 'cape town', 'durban', 'port elizabeth', 
                 'east london', 'bloemfontein', 'nelspruit', 'polokwane', 'george', 
                 'plettenberg bay', 'east london']
        
        text_lower = text.lower()
        for city in cities:
            if city in text_lower:
                return city.title()
        return "South Africa"
    
    def _extract_reference(self, text: str) -> Optional[str]:
        patterns = [
            r'RFQ\s*[No.]*\s*(\d{6,12})',
            r'PR\s*(\d{6,12})',
            r'TENDER\s*NO[.:]\s*(\w{5,20})'
        ]
        for pattern in patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                return match.group(0).replace(' ', '').replace(':', '')[:30]
        return None
    
    def _extract_closing_date(self, text: str) -> str:
        match = re.search(r'(\d{1,2}[\/\-.]\d{1,2}[\/\-.]\d{2,4})', text)
        if match:
            return match.group(1)
        return "TBC"
    
    def _extract_scope_description(self, text: str, sector: str) -> str:
        lines = text.split('\n')
        for line in lines:
            clean = line.strip()
            if len(clean) > 50 and any(word in clean.lower() for word in ['scope', 'description', 'background']):
                return clean[:200]
        return f"{sector.title()} services as per tender"
    
    def _extract_client_type(self, text: str) -> str:
        if 'MUNICIPALITY' in text:
            return 'Municipality'
        elif 'GOVERNMENT' in text:
            return 'Government'
        elif 'SCHOOL' in text:
            return 'Educational'
        return 'Public Entity'
    
    def _extract_client_name(self, text: str) -> str:
        match = re.search(r'(?:FROM|BY)\s*([A-Z][A-Za-z\s]{10,50})', text)
        if match:
            return match.group(1).strip()[:50]
        return "Client"
    
    def _calculate_overall_confidence(self, worker_data, duration_data, scope_data) -> str:
        scores = []
        if not worker_data.get('estimated'):
            scores.append(1.0)
        else:
            scores.append(0.5 if worker_data.get('confidence') == 'Medium' else 0.3)
        if not duration_data.get('estimated'):
            scores.append(1.0)
        else:
            scores.append(0.7)
        if not scope_data.get('estimated'):
            scores.append(1.0)
        else:
            scores.append(0.6)
        
        avg = sum(scores) / len(scores)
        if avg >= 0.8:
            return "High"
        elif avg >= 0.5:
            return "Medium"
        else:
            return "Low - Verify"
    
    async def _extract_pdf(self, pdf_path: Path) -> str:
        text = ""
        try:
            with pdfplumber.open(pdf_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
        except Exception as e:
            print(f"pdfplumber error: {e}")
        if len(text.strip()) < 500:
            text = await self._ocr_pdf(pdf_path)
        return text
    
    async def _ocr_pdf(self, pdf_path: Path) -> str:
        text = ""
        try:
            images = convert_from_path(str(pdf_path), dpi=150)
            for i, image in enumerate(images[:20]):  # INCREASED to 20 pages
                page_text = pytesseract.image_to_string(image)
                text += f"\nPAGE {i+1}:\n{page_text}\n"
        except Exception as e:
            print(f"OCR error: {e}")
        return text
    
    async def _extract_docx(self, docx_path: Path) -> str:
        if not DOCX_AVAILABLE:
            try:
                import zipfile
                with zipfile.ZipFile(docx_path) as zf:
                    xml_content = zf.read('word/document.xml')
                    text = re.sub(r'<[^>]+>', '', xml_content.decode('utf-8', errors='ignore'))
                    return text
            except:
                return ""
        try:
            doc = Document(str(docx_path))
            full_text = []
            for para in doc.paragraphs:
                if para.text.strip():
                    full_text.append(para.text)
            for table in doc.tables:
                for row in table.rows:
                    row_text = ' | '.join([cell.text.strip() for cell in row.cells])
                    if row_text:
                        full_text.append(row_text)
            return '\n'.join(full_text)
        except:
            return ""
    
    async def _extract_text(self, txt_path: Path) -> str:
        try:
            with open(txt_path, 'r', encoding='utf-8', errors='ignore') as f:
                return f.read()
        except:
            return ""

# Backward compatibility
class PDFProcessor(DocumentExtractor):
    pass
