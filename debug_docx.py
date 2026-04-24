# debug_docx.py
from docx import Document
doc = Document("workspace/tenders/1/RFQ_Document_PR10115496.docx")
print("PARAGRAPHS FOUND:", len(doc.paragraphs))
for i, para in enumerate(doc.paragraphs[:10]):
    print(f"{i}: {para.text[:100]}")
